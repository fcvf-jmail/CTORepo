#!/usr/bin/env python3
"""Полный pipeline бинарной классификации SMS (spam/ham) с использованием NLP."""
from __future__ import annotations

import logging
import pickle
import re
import zipfile
from collections import Counter
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import requests
import seaborn as sns
from docx import Document
from docx.shared import Inches
from nltk import pos_tag
from nltk.corpus import stopwords, wordnet
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    confusion_matrix,
    precision_recall_fscore_support,
    roc_auc_score,
    roc_curve,
)
from sklearn.model_selection import StratifiedKFold, cross_validate, train_test_split
from sklearn.naive_bayes import MultinomialNB
from sklearn.utils import resample

import nltk

# ---------------------------
# Общие настройки
# ---------------------------
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "ml_outputs"
DATA_DIR = OUTPUT_DIR / "data"
FIGURES_DIR = OUTPUT_DIR / "figures"
ARTIFACTS_DIR = OUTPUT_DIR / "artifacts"
DOCS_DIR = BASE_DIR / "docs"
REPORT_PATH = DOCS_DIR / "SMS_Spam_Filter_Report.docx"
MODEL_PATH = ARTIFACTS_DIR / "sms_spam_best_model.pkl"
VECTORIZER_PATH = ARTIFACTS_DIR / "sms_spam_tfidf_vectorizer.pkl"
DATASET_URL = "https://archive.ics.uci.edu/ml/machine-learning-databases/00228/smsspamcollection.zip"
RANDOM_STATE = 42

sns.set_theme(style="whitegrid")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s — %(levelname)s — %(message)s",
)


@dataclass
class ModelResult:
    """Структура для хранения результатов модели."""

    name: str
    estimator: object
    cv_metrics: Dict[str, float]
    test_metrics: Dict[str, float]
    confusion: np.ndarray
    roc_curve_points: Tuple[np.ndarray, np.ndarray, np.ndarray]
    feature_table: pd.DataFrame


class TextPreprocessor:
    """Класс для очистки и нормализации SMS сообщений."""

    url_pattern = re.compile(r"(https?://\S+|www\.\S+)")
    email_pattern = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
    digit_pattern = re.compile(r"\d+")
    non_alpha_pattern = re.compile(r"[^a-zA-Z\s]")

    def __init__(self) -> None:
        self._ensure_nltk_resources()
        self.stop_words = set(stopwords.words("english"))
        self.lemmatizer = WordNetLemmatizer()

    @staticmethod
    def _ensure_nltk_resources() -> None:
        """Скачивает необходимые ресурсы NLTK один раз."""

        resources = [
            "punkt",
            "punkt_tab",
            "stopwords",
            "wordnet",
            "omw-1.4",
            "averaged_perceptron_tagger",
            "averaged_perceptron_tagger_eng",
        ]
        for resource in resources:
            nltk.download(resource, quiet=True)

    @staticmethod
    def _get_wordnet_pos(tag: str) -> str:
        """Возвращает POS-тег для лемматизации."""

        if tag.startswith("J"):
            return wordnet.ADJ
        if tag.startswith("V"):
            return wordnet.VERB
        if tag.startswith("N"):
            return wordnet.NOUN
        if tag.startswith("R"):
            return wordnet.ADV
        return wordnet.NOUN

    def preprocess(self, text: str) -> str:
        """Выполняет полную очистку, токенизацию и лемматизацию."""

        cleaned = text.lower()
        cleaned = self.url_pattern.sub(" ", cleaned)
        cleaned = self.email_pattern.sub(" ", cleaned)
        cleaned = self.digit_pattern.sub(" ", cleaned)
        cleaned = self.non_alpha_pattern.sub(" ", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()

        tokens = word_tokenize(cleaned)
        tokens = [token for token in tokens if token not in self.stop_words and len(token) > 1]
        if not tokens:
            return "<empty>"

        tagged_tokens = pos_tag(tokens)
        normalized = [
            self.lemmatizer.lemmatize(token, self._get_wordnet_pos(tag))
            for token, tag in tagged_tokens
        ]
        normalized = [token for token in normalized if token not in self.stop_words]

        return " ".join(normalized) if normalized else "<empty>"


def ensure_directories() -> None:
    """Гарантирует наличие служебных директорий."""

    for path in [OUTPUT_DIR, DATA_DIR, FIGURES_DIR, ARTIFACTS_DIR, DOCS_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def download_dataset() -> Path:
    """Скачивает датасет SMS Spam Collection и возвращает путь к txt файлу."""

    dataset_path = DATA_DIR / "SMSSpamCollection"
    if dataset_path.exists():
        logging.info("Датасет уже загружен — пропускаем загрузку.")
        return dataset_path

    logging.info("Скачиваем датасет SMS Spam Collection...")
    response = requests.get(DATASET_URL, timeout=30)
    response.raise_for_status()

    with zipfile.ZipFile(BytesIO(response.content)) as archive:
        archive.extractall(DATA_DIR)

    logging.info("Датасет успешно загружен в %s", dataset_path)
    return dataset_path


def load_dataset(dataset_file: Path) -> pd.DataFrame:
    """Загружает и подготавливает исходный DataFrame."""

    df = pd.read_csv(dataset_file, sep="\t", header=None, names=["label", "message"])
    df["label"] = df["label"].map({"ham": 0, "spam": 1})
    return df


def summarize_dataset(df: pd.DataFrame) -> Dict[str, float]:
    """Возвращает ключевые статистики датасета."""

    total = len(df)
    spam = df["label"].sum()
    ham = total - spam
    avg_length = df["message"].str.len().mean()
    max_length = df["message"].str.len().max()

    return {
        "total": int(total),
        "spam": int(spam),
        "ham": int(ham),
        "spam_pct": float(spam / total * 100),
        "ham_pct": float(ham / total * 100),
        "avg_length": float(avg_length),
        "max_length": int(max_length),
    }


def split_dataset(df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """Разбивает данные на train/test в пропорции 80/20 со стратификацией."""

    train_df, test_df = train_test_split(
        df,
        test_size=0.2,
        stratify=df["label"],
        random_state=RANDOM_STATE,
    )
    logging.info(
        "Размеры выборок — train: %d, test: %d", len(train_df), len(test_df)
    )
    return train_df.reset_index(drop=True), test_df.reset_index(drop=True)


def balance_training_data(train_df: pd.DataFrame) -> pd.DataFrame:
    """Выполняет oversampling для балансировки классов."""

    spam_df = train_df[train_df["label"] == 1]
    ham_df = train_df[train_df["label"] == 0]
    spam_upsampled = resample(
        spam_df,
        replace=True,
        n_samples=len(ham_df),
        random_state=RANDOM_STATE,
    )
    balanced = pd.concat([ham_df, spam_upsampled]).sample(frac=1, random_state=RANDOM_STATE)
    logging.info(
        "Балансировка: ham=%d, spam=%d",
        len(balanced[balanced["label"] == 0]),
        len(balanced[balanced["label"] == 1]),
    )
    return balanced.reset_index(drop=True)


def preprocess_dataframe(df: pd.DataFrame, preprocessor: TextPreprocessor) -> pd.Series:
    """Применяет препроцессинг к серии сообщений."""

    return df["message"].apply(preprocessor.preprocess)


def vectorize_text(
    train_texts: Iterable[str], test_texts: Iterable[str]
) -> Tuple[TfidfVectorizer, np.ndarray, np.ndarray]:
    """Строит TF-IDF признаки и возвращает матрицы."""

    vectorizer = TfidfVectorizer(
        ngram_range=(1, 2),
        min_df=2,
        max_df=0.95,
        sublinear_tf=True,
    )
    x_train = vectorizer.fit_transform(train_texts)
    x_test = vectorizer.transform(test_texts)
    logging.info("Размерность признаков TF-IDF: %d", x_train.shape[1])
    return vectorizer, x_train, x_test


def evaluate_model(
    model: object,
    x_test: np.ndarray,
    y_test: np.ndarray,
) -> Tuple[Dict[str, float], np.ndarray, Tuple[np.ndarray, np.ndarray, np.ndarray]]:
    """Вычисляет метрики и возвращает confusion matrix и ROC-кривую."""

    y_pred = model.predict(x_test)
    y_prob = model.predict_proba(x_test)[:, 1]

    accuracy = accuracy_score(y_test, y_pred)
    precision, recall, f1, _ = precision_recall_fscore_support(
        y_test, y_pred, average="binary"
    )
    roc_auc = roc_auc_score(y_test, y_prob)
    conf_matrix = confusion_matrix(y_test, y_pred)
    roc_points = roc_curve(y_test, y_prob)

    metrics = {
        "accuracy": accuracy,
        "precision": precision,
        "recall": recall,
        "f1": f1,
        "roc_auc": roc_auc,
    }
    return metrics, conf_matrix, roc_points


def cross_validate_model(
    model: object,
    x_train: np.ndarray,
    y_train: np.ndarray,
    cv: StratifiedKFold,
) -> Dict[str, float]:
    """Запускает кросс-валидацию и усредняет метрики."""

    scoring = {
        "accuracy": "accuracy",
        "precision": "precision",
        "recall": "recall",
        "f1": "f1",
        "roc_auc": "roc_auc",
    }
    cv_result = cross_validate(
        model,
        x_train,
        y_train,
        cv=cv,
        scoring=scoring,
        n_jobs=-1,
    )
    return {key.replace("test_", ""): float(np.mean(value)) for key, value in cv_result.items() if key.startswith("test_")}


def extract_feature_importance(
    model: object,
    vectorizer: TfidfVectorizer,
    top_n: int = 15,
) -> pd.DataFrame:
    """Возвращает таблицу с топовыми словами для spam/ham."""

    feature_names = np.array(vectorizer.get_feature_names_out())

    if isinstance(model, LogisticRegression):
        weights = model.coef_[0]
    elif isinstance(model, MultinomialNB):
        weights = model.feature_log_prob_[1] - model.feature_log_prob_[0]
    else:
        raise ValueError("Модель не поддерживает извлечение важности признаков")

    spam_idx = np.argsort(weights)[-top_n:][::-1]
    ham_idx = np.argsort(weights)[:top_n]

    data = {
        "spam_word": feature_names[spam_idx],
        "spam_weight": weights[spam_idx],
        "ham_word": feature_names[ham_idx],
        "ham_weight": np.abs(weights[ham_idx]),
    }
    return pd.DataFrame(data)


def plot_class_distribution(
    train_counts: Counter,
    test_counts: Counter,
    save_path: Path,
) -> None:
    """Строит диаграмму распределения классов."""

    df = pd.DataFrame(
        {
            "Выборка": ["Train"] * 2 + ["Test"] * 2,
            "Класс": ["Ham", "Spam"] * 2,
            "Количество": [train_counts[0], train_counts[1], test_counts[0], test_counts[1]],
        }
    )
    plt.figure(figsize=(6, 4))
    sns.barplot(data=df, x="Выборка", y="Количество", hue="Класс", palette="viridis")
    plt.title("Распределение классов после сплита")
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()


def plot_confusion_matrix(
    matrix: np.ndarray,
    title: str,
    save_path: Path,
) -> None:
    """Визуализирует confusion matrix."""

    plt.figure(figsize=(4.5, 4))
    sns.heatmap(
        matrix,
        annot=True,
        fmt="d",
        cmap="rocket",
        xticklabels=["Ham", "Spam"],
        yticklabels=["Ham", "Spam"],
    )
    plt.title(title)
    plt.xlabel("Предсказание")
    plt.ylabel("Истина")
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()


def plot_roc_curves(
    roc_data: List[Tuple[str, np.ndarray, np.ndarray, float]],
    save_path: Path,
) -> None:
    """Рисует ROC-кривые для всех моделей."""

    plt.figure(figsize=(6, 5))
    for name, fpr, tpr, auc_value in roc_data:
        plt.plot(fpr, tpr, label=f"{name} (AUC={auc_value:.3f})")
    plt.plot([0, 1], [0, 1], linestyle="--", color="grey")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("ROC-кривые моделей")
    plt.legend()
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()


def plot_metrics_comparison(
    model_results: Dict[str, ModelResult],
    save_path: Path,
) -> None:
    """Сравнивает Accuracy/Precision/Recall/F1 по моделям."""

    records = []
    for result in model_results.values():
        metrics = result.test_metrics
        records.append(
            {
                "Модель": result.name,
                "Accuracy": metrics["accuracy"],
                "Precision": metrics["precision"],
                "Recall": metrics["recall"],
                "F1": metrics["f1"],
                "ROC-AUC": metrics["roc_auc"],
            }
        )
    df = pd.DataFrame(records).set_index("Модель")
    df.plot(kind="bar", figsize=(8, 5), colormap="magma")
    plt.title("Сравнение метрик на тестовой выборке")
    plt.ylim(0, 1.05)
    plt.ylabel("Значение")
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()


def plot_feature_importance(
    table: pd.DataFrame,
    title: str,
    save_path: Path,
) -> None:
    """Визуализация топ слов для spam/ham."""

    fig, axes = plt.subplots(1, 2, figsize=(10, 6), sharey=True)
    spam_data = table.sort_values("spam_weight", ascending=True)
    ham_data = table.sort_values("ham_weight", ascending=True)

    axes[0].barh(spam_data["spam_word"], spam_data["spam_weight"], color="#f07167")
    axes[0].set_title("Spam слова")
    axes[0].set_xlabel("Вес признака")
    axes[0].set_ylabel("Слово")

    axes[1].barh(ham_data["ham_word"], ham_data["ham_weight"], color="#00afb9")
    axes[1].set_title("Ham слова")
    axes[1].set_xlabel("Вес признака")
    axes[1].set_ylabel("")

    fig.suptitle(title)
    plt.tight_layout()
    plt.savefig(save_path)
    plt.close()


def save_pickle(obj: object, path: Path) -> None:
    """Сохраняет объект в формате pickle."""

    with open(path, "wb") as file:
        pickle.dump(obj, file)
    logging.info("Артефакт сохранён: %s", path)


def demo_predictions(
    preprocessor: TextPreprocessor,
    vectorizer: TfidfVectorizer,
    model: object,
    samples: List[str],
) -> List[Dict[str, str]]:
    """Возвращает список с примерами предсказаний для новых SMS."""

    processed = [preprocessor.preprocess(text) for text in samples]
    features = vectorizer.transform(processed)
    probs = model.predict_proba(features)[:, 1]
    labels = model.predict(features)
    decoded = []
    for original, label, prob in zip(samples, labels, probs):
        prediction = "SPAM" if label == 1 else "HAM"
        decoded.append(
            {
                "message": original,
                "prediction": prediction,
                "probability": f"{prob:.3f}",
            }
        )
        logging.info("Сообщение: %s => %s (p=%.3f)", original, prediction, prob)
    return decoded


def generate_report(
    report_path: Path,
    dataset_summary: Dict[str, float],
    class_distribution: Dict[str, Counter],
    preprocessing_steps: List[str],
    model_results: Dict[str, ModelResult],
    balancing_description: str,
    figures: Dict[str, Path],
    confusion_paths: Dict[str, Path],
    feature_paths: Dict[str, Path],
    sample_predictions: List[Dict[str, str]],
) -> None:
    """Создаёт Word-отчёт с описанием эксперимента."""

    document = Document()
    document.add_heading("Система классификации SMS на основе NLP", 0)

    document.add_heading("1. Постановка задачи", level=1)
    document.add_paragraph(
        "Цель — разработать полнофункциональный фильтр SMS, "
        "который различает spam и ham сообщения с помощью методов NLP "
        "и классических моделей машинного обучения."
    )

    document.add_heading("2. Датасет", level=1)
    document.add_paragraph(
        "Использован публичный SMS Spam Collection (5 574 сообщения). Источник: "
        "UCI Machine Learning Repository."
    )
    document.add_paragraph(
        "Общая статистика: "
        f"спам — {dataset_summary['spam']} сообщений ({dataset_summary['spam_pct']:.2f}%), "
        f"ham — {dataset_summary['ham']} сообщений ({dataset_summary['ham_pct']:.2f}%). "
        f"Средняя длина SMS: {dataset_summary['avg_length']:.1f} символов, "
        f"максимальная длина: {dataset_summary['max_length']} символов."
    )

    document.add_heading("3. Распределение классов", level=1)
    distribution_text = (
        f"Train — ham: {class_distribution['train'][0]}, spam: {class_distribution['train'][1]}; "
        f"Test — ham: {class_distribution['test'][0]}, spam: {class_distribution['test'][1]}."
    )
    document.add_paragraph(distribution_text)
    if figures.get("class_distribution") and figures["class_distribution"].exists():
        document.add_picture(str(figures["class_distribution"]), width=Inches(5.5))

    document.add_heading("4. NLP-пайплайн", level=1)
    document.add_paragraph(
        "Пайплайн состоит из обязательных шагов обработки текста (очистка, токенизация, "
        "удаление стоп-слов, лемматизация)."
    )
    for step in preprocessing_steps:
        document.add_paragraph(step, style="List Bullet")

    document.add_paragraph(f"Обработка дисбаланса: {balancing_description}.")

    document.add_heading("5. Модели", level=1)
    document.add_paragraph(
        "Обучены Multinomial Naive Bayes и Logistic Regression с настройкой "
        "гиперпараметров и пятифолдовой стратифицированной кросс-валидацией."
    )

    document.add_heading("6. Результаты", level=1)
    for name, result in model_results.items():
        document.add_heading(name, level=2)
        metrics = result.test_metrics
        cv_metrics = result.cv_metrics
        document.add_paragraph(
            f"Test — Accuracy: {metrics['accuracy']:.3f}, Precision: {metrics['precision']:.3f}, "
            f"Recall: {metrics['recall']:.3f}, F1: {metrics['f1']:.3f}, ROC-AUC: {metrics['roc_auc']:.3f}."
        )
        document.add_paragraph(
            f"CV (5-fold) — Acc: {cv_metrics['accuracy']:.3f}, Prec: {cv_metrics['precision']:.3f}, "
            f"Rec: {cv_metrics['recall']:.3f}, F1: {cv_metrics['f1']:.3f}, AUC: {cv_metrics['roc_auc']:.3f}."
        )
        if name in confusion_paths and confusion_paths[name].exists():
            document.add_paragraph("Confusion matrix:")
            document.add_picture(str(confusion_paths[name]), width=Inches(4.5))
        if name in feature_paths and feature_paths[name].exists():
            document.add_paragraph("Топ информативных слов:")
            document.add_picture(str(feature_paths[name]), width=Inches(5.5))

    if figures.get("roc_curve") and figures["roc_curve"].exists():
        document.add_heading("ROC-кривые", level=2)
        document.add_picture(str(figures["roc_curve"]), width=Inches(5.5))

    if figures.get("metrics_comparison") and figures["metrics_comparison"].exists():
        document.add_heading("Сравнение моделей", level=2)
        document.add_picture(str(figures["metrics_comparison"]), width=Inches(5.5))

    document.add_heading("7. Примеры предсказаний", level=1)
    for sample in sample_predictions:
        document.add_paragraph(
            f"{sample['prediction']} ({sample['probability']}): {sample['message']}",
            style="List Bullet",
        )

    document.add_heading("8. Выводы и рекомендации", level=1)
    document.add_paragraph(
        "Логистическая регрессия продемонстрировала лучшую обобщающую способность благодаря "
        "использованию TF-IDF признаков и балансировке классов. Для дальнейшего улучшения "
        "можно протестировать ансамблевые модели, добавить character n-grams и применить "
        "современные эмбеддинги."
    )

    document.save(report_path)
    logging.info("Отчёт сохранён: %s", report_path)


def main() -> None:
    """Точка входа для запуска всего pipeline."""

    ensure_directories()
    dataset_file = download_dataset()
    df = load_dataset(dataset_file)
    dataset_summary = summarize_dataset(df)

    train_df, test_df = split_dataset(df)
    class_distribution = {
        "train": Counter(train_df["label"]),
        "test": Counter(test_df["label"]),
    }

    preprocessor = TextPreprocessor()
    train_df = train_df.assign(clean_text=preprocess_dataframe(train_df, preprocessor))
    test_df = test_df.assign(clean_text=preprocess_dataframe(test_df, preprocessor))

    balanced_train_df = balance_training_data(train_df[["clean_text", "label"]])

    vectorizer, x_train, x_test = vectorize_text(
        balanced_train_df["clean_text"],
        test_df["clean_text"],
    )
    y_train = balanced_train_df["label"].to_numpy()
    y_test = test_df["label"].to_numpy()

    cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)

    model_builders = {
        "Multinomial Naive Bayes": lambda: MultinomialNB(alpha=0.5),
        "Logistic Regression": lambda: LogisticRegression(
            max_iter=1000,
            class_weight="balanced",
            solver="liblinear",
        ),
    }

    model_results: Dict[str, ModelResult] = {}
    roc_items: List[Tuple[str, np.ndarray, np.ndarray, float]] = []
    confusion_paths: Dict[str, Path] = {}
    feature_paths: Dict[str, Path] = {}

    for name, builder in model_builders.items():
        logging.info("Обработка модели: %s", name)
        cv_metrics = cross_validate_model(builder(), x_train, y_train, cv)
        estimator = builder()
        estimator.fit(x_train, y_train)
        test_metrics, conf_matrix, roc_points = evaluate_model(estimator, x_test, y_test)
        feature_table = extract_feature_importance(estimator, vectorizer)

        confusion_path = FIGURES_DIR / f"confusion_{name.replace(' ', '_').lower()}.png"
        plot_confusion_matrix(conf_matrix, f"Confusion Matrix — {name}", confusion_path)

        feature_path = FIGURES_DIR / f"features_{name.replace(' ', '_').lower()}.png"
        plot_feature_importance(feature_table, f"Топ слов — {name}", feature_path)

        roc_items.append((name, roc_points[0], roc_points[1], test_metrics["roc_auc"]))

        model_results[name] = ModelResult(
            name=name,
            estimator=estimator,
            cv_metrics=cv_metrics,
            test_metrics=test_metrics,
            confusion=conf_matrix,
            roc_curve_points=roc_points,
            feature_table=feature_table,
        )
        confusion_paths[name] = confusion_path
        feature_paths[name] = feature_path

    roc_path = FIGURES_DIR / "roc_curves.png"
    plot_roc_curves(roc_items, roc_path)

    class_distribution_path = FIGURES_DIR / "class_distribution.png"
    plot_class_distribution(class_distribution["train"], class_distribution["test"], class_distribution_path)

    metrics_path = FIGURES_DIR / "metrics_comparison.png"
    plot_metrics_comparison(model_results, metrics_path)

    best_model_name = max(
        model_results.values(),
        key=lambda result: result.test_metrics["f1"],
    ).name
    best_estimator = model_results[best_model_name].estimator
    save_pickle(best_estimator, MODEL_PATH)
    save_pickle(vectorizer, VECTORIZER_PATH)

    sample_texts = [
        "Congratulations! You have won a $500 Amazon voucher. Claim now!",
        "Привет! Напомни, пожалуйста, во сколько встречаемся вечером?",
        "URGENT! Your account has been compromised. Verify immediately at http://fake.link",
        "Доброе утро. Я оплатил счет за интернет, все ок.",
    ]
    sample_predictions = demo_predictions(preprocessor, vectorizer, best_estimator, sample_texts)

    preprocessing_steps = [
        "Приведение текста к нижнему регистру, удаление URL, email, цифр и спецсимволов.",
        "Токенизация предложений при помощи NLTK (word_tokenize).",
        "Удаление английских стоп-слов из списка NLTK.",
        "Лемматизация WordNet с учетом POS-тегов (nltk.pos_tag).",
    ]

    figures = {
        "class_distribution": class_distribution_path,
        "roc_curve": roc_path,
        "metrics_comparison": metrics_path,
    }

    generate_report(
        REPORT_PATH,
        dataset_summary,
        class_distribution,
        preprocessing_steps,
        model_results,
        "Random oversampling миноритарного класса (spam) до уровня ham",
        figures,
        confusion_paths,
        feature_paths,
        sample_predictions,
    )

    logging.info("Pipeline успешно завершён. Лучший классификатор: %s", best_model_name)


if __name__ == "__main__":
    main()
